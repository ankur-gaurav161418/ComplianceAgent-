"""
app.py
------
Streamlit chat app that answers questions from your PDF documents.

How it works:
  1. Loads the FAISS vector database built by create_vector_db.py
  2. When you ask a question, finds the most relevant text chunks
  3. Sends those chunks + conversation history to Ollama (llama3.1:8b)
  4. Streams the answer back and shows which PDFs it came from

Requirements:
  pip install streamlit sentence-transformers faiss-cpu numpy
  # Ollama must be running locally: https://ollama.com
  # Pull the model first: ollama pull llama3.1:8b

Run:
  streamlit run app.py
"""

import os
import re
import io
import json
import base64
import hashlib
import datetime
from html import escape as esc

import requests
import numpy as np
import faiss
import fitz
import streamlit as st
from sentence_transformers import SentenceTransformer

import letterhead

# --------------------------------------------------------------------------
# Configuration
# --------------------------------------------------------------------------
VECTORSTORE_FOLDER = "vectorstore"
INDEX_FILE = os.path.join(VECTORSTORE_FOLDER, "index.faiss")
CHUNKS_FILE = os.path.join(VECTORSTORE_FOLDER, "chunks.json")
EMBEDDING_MODEL = "all-MiniLM-L6-v2"
TOP_K = 5
OLLAMA_BASE_URL = "http://localhost:11434"
OLLAMA_CHAT_URL = OLLAMA_BASE_URL + "/api/chat"
OLLAMA_MODEL = "llama3.1:8b"
MAX_HISTORY = 10

APP_NAME = "Legal Compliance AI Agent"
APP_TAGLINE = "Enterprise document intelligence for legal & compliance teams"

ASSETS_FOLDER = "assets"
FAVICON_FILE = os.path.join(ASSETS_FOLDER, "favicon.png")
LOGO_ICON_FILE = os.path.join(ASSETS_FOLDER, "logo-icon-128.png")

BLOCKED_TERMS = ["hack", "fraud", "tax evasion", "steal"]
GUARDRAIL_MESSAGE = (
    "This request was blocked by compliance guardrails because it matched a "
    "restricted term (“{}”). Please rephrase your question."
)

EXAMPLE_QUESTIONS = [
    "What personal data does GDPR protect?",
    "What are the penalties under the DPDP Act, 2023?",
    "Who is a Data Fiduciary under Indian law?",
]

# --------------------------------------------------------------------------
# Attachments (upload files as input: PDF, screenshots/images, Word, text)
# --------------------------------------------------------------------------
PDF_EXTS = {"pdf"}
IMAGE_EXTS = {"png", "jpg", "jpeg", "webp", "bmp", "tif", "tiff", "gif"}
DOCX_EXTS = {"docx"}
TEXT_EXTS = {
    "txt", "md", "markdown", "csv", "tsv", "json", "log",
    "xml", "html", "htm", "yaml", "yml", "ini", "cfg", "py", "js", "ts",
}
# Formats offered in the picker but not directly parseable -> user is told to convert.
CONVERT_EXTS = {"doc", "odt", "rtf", "pages"}
SUPPORTED_UPLOAD_EXTS = sorted(
    PDF_EXTS | IMAGE_EXTS | DOCX_EXTS | TEXT_EXTS | CONVERT_EXTS
)

ATTACH_CHUNK_SIZE = 800
ATTACH_CHUNK_OVERLAP = 100
ATTACH_TOP_K = 5
DEFAULT_ATTACH_PROMPT = (
    "Summarise the attached file(s) and highlight the key points that are "
    "relevant to legal compliance."
)


# --------------------------------------------------------------------------
# Guardrails
# --------------------------------------------------------------------------
def check_guardrails(text):
    lowered = text.lower()
    for term in BLOCKED_TERMS:
        if " " in term:
            if term in lowered:
                return True, term
        else:
            if re.search(r"\b" + re.escape(term) + r"\b", lowered):
                return True, term
    return False, None


# --------------------------------------------------------------------------
# Data / model loading
# --------------------------------------------------------------------------
@st.cache_resource(show_spinner=False)
def load_vectorstore():
    if not os.path.exists(INDEX_FILE) or not os.path.exists(CHUNKS_FILE):
        return None, None
    index = faiss.read_index(INDEX_FILE)
    with open(CHUNKS_FILE, "r", encoding="utf-8") as f:
        chunks = json.load(f)
    return index, chunks


@st.cache_resource(show_spinner=False)
def load_embedding_model():
    return SentenceTransformer(EMBEDDING_MODEL)


@st.cache_data(show_spinner=False)
def load_logo_base64():
    if not os.path.exists(LOGO_ICON_FILE):
        return None
    with open(LOGO_ICON_FILE, "rb") as f:
        return base64.b64encode(f.read()).decode()


@st.cache_data(ttl=10, show_spinner=False)
def check_ollama_status():
    try:
        r = requests.get(OLLAMA_BASE_URL + "/api/tags", timeout=2)
        r.raise_for_status()
        names = [m.get("name", "") for m in r.json().get("models", [])]
        model_ready = any(n == OLLAMA_MODEL or n.startswith(OLLAMA_MODEL) for n in names)
        return True, model_ready
    except requests.exceptions.RequestException:
        return False, False


def embed_query(model, query):
    return model.encode([query], convert_to_numpy=True).astype(np.float32)


def search_kb(query_vec, index, chunks, top_k):
    """Retrieve the most relevant chunks from the persistent knowledge base."""
    if index is None or not chunks:
        return []
    distances, indices = index.search(query_vec, top_k)
    results = []
    for dist, idx in zip(distances[0], indices[0]):
        if idx == -1:
            continue
        chunk = chunks[idx]
        results.append({
            "text": chunk["text"],
            "source": chunk["source"],
            "chunk_id": chunk["chunk_id"],
            "score": float(dist),
            "origin": "kb",
        })
    return results


def search_attachments(query_vec, top_k):
    """Retrieve the most relevant chunks from files attached this session."""
    attachments = st.session_state.get("attachments", [])
    q = query_vec[0]
    hits = []
    for att in attachments:
        embeddings = att.get("embeddings")
        if embeddings is None or len(att["chunks"]) == 0:
            continue
        distances = np.linalg.norm(embeddings - q, axis=1)
        for i, dist in enumerate(distances):
            hits.append({
                "text": att["chunks"][i],
                "source": "📎 " + att["name"],
                "chunk_id": i,
                "score": float(dist),
                "origin": "attachment",
            })
    hits.sort(key=lambda h: h["score"])
    return hits[:top_k]


# --------------------------------------------------------------------------
# Attachment text extraction
# --------------------------------------------------------------------------
def file_ext(name):
    return name.rsplit(".", 1)[-1].lower() if "." in name else ""


@st.cache_resource(show_spinner=False)
def load_ocr_engine():
    from rapidocr_onnxruntime import RapidOCR
    return RapidOCR()


def _looks_binary(data):
    sample = data[:2048]
    if b"\x00" in sample:
        return True
    # Zip-based office formats (docx handled separately, but doc/odt/xlsx/pptx are not)
    if sample[:2] == b"PK":
        return True
    return False


def _extract_pdf(data):
    doc = fitz.open(stream=data, filetype="pdf")
    pages = [doc[i].get_text() for i in range(len(doc))]
    doc.close()
    return "\n".join(pages).strip()


def _extract_docx(data):
    from docx import Document
    document = Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(p for p in parts if p).strip()


def _extract_image(data):
    engine = load_ocr_engine()
    result, _ = engine(data)
    if not result:
        return ""
    return "\n".join(line[1] for line in result).strip()


def _extract_plain(data):
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(enc).strip()
        except (UnicodeDecodeError, UnicodeError):
            continue
    return data.decode("utf-8", errors="ignore").strip()


def extract_text_from_upload(name, data):
    """Return (text, error). error is a human-readable string or None."""
    ext = file_ext(name)
    try:
        if ext in PDF_EXTS:
            return _extract_pdf(data), None
        if ext in IMAGE_EXTS:
            return _extract_image(data), None
        if ext in DOCX_EXTS:
            return _extract_docx(data), None
        if ext in CONVERT_EXTS or _looks_binary(data):
            return "", ("Can't read this format directly — please convert "
                        "it to PDF or .docx and re-attach.")
        return _extract_plain(data), None
    except Exception as exc:  # noqa: BLE001 - surface any parser error to the user
        return "", "Could not read file: {}".format(exc)


def chunk_text(text, size, overlap):
    chunks = []
    start = 0
    n = len(text)
    while start < n:
        chunk = text[start:start + size].strip()
        if chunk:
            chunks.append(chunk)
        start += max(1, size - overlap)
    return chunks


def process_attachments(files, embedding_model):
    """Extract, chunk and embed newly attached files into session state.

    Returns a list of (name, status, detail) notes for the caller to display.
    """
    if "attachments" not in st.session_state:
        st.session_state.attachments = []
    seen = {att["hash"] for att in st.session_state.attachments}
    notes = []
    for f in files:
        data = f.getvalue()
        digest = hashlib.md5(data).hexdigest()
        if digest in seen:
            notes.append((f.name, "duplicate", "already attached"))
            continue
        seen.add(digest)
        text, error = extract_text_from_upload(f.name, data)
        if error:
            notes.append((f.name, "error", error))
            continue
        text_chunks = chunk_text(text, ATTACH_CHUNK_SIZE, ATTACH_CHUNK_OVERLAP)
        embeddings = None
        if text_chunks:
            embeddings = embedding_model.encode(
                text_chunks, convert_to_numpy=True).astype(np.float32)
        st.session_state.attachments.append({
            "name": f.name,
            "hash": digest,
            "kind": file_ext(f.name) or "file",
            "char_count": len(text),
            "chunks": text_chunks,
            "embeddings": embeddings,
        })
        if text_chunks:
            notes.append((f.name, "ok", "{} chars, {} chunks".format(
                len(text), len(text_chunks))))
        else:
            notes.append((f.name, "empty", "no readable text found"))
    return notes


def build_prompt_messages(history, context_text, user_question):
    system_message = {
        "role": "system",
        "content": (
            "You are a helpful assistant that answers questions based on the"
            " provided document excerpts. Use the excerpts below as your primary"
            " source of information. If the answer is not in the excerpts, say"
            " so clearly.\n\nDOCUMENT EXCERPTS:\n" + context_text
        ),
    }
    recent_history = history[-(MAX_HISTORY * 2):]
    return [system_message] + recent_history + [{"role": "user", "content": user_question}]


def stream_ollama(messages):
    payload = {"model": OLLAMA_MODEL, "messages": messages, "stream": True}
    try:
        response = requests.post(OLLAMA_CHAT_URL, json=payload, stream=True, timeout=120)
        response.raise_for_status()
    except requests.exceptions.ConnectionError:
        yield "ERROR: Cannot connect to Ollama. Run 'ollama serve' and pull the model."
        return
    except requests.exceptions.RequestException as e:
        yield "ERROR: {}".format(e)
        return
    for line in response.iter_lines():
        if not line:
            continue
        try:
            data = json.loads(line)
            token = data.get("message", {}).get("content", "")
            if token:
                yield token
            if data.get("done", False):
                break
        except json.JSONDecodeError:
            continue


# --------------------------------------------------------------------------
# Design system (CSS)
# --------------------------------------------------------------------------
def inject_css():
    st.markdown("""
    <style>
    :root {
        --ca-bg: #F6F7FA;
        --ca-surface: #FFFFFF;
        --ca-surface-alt: #F0F2F6;
        --ca-border: #E3E6ED;
        --ca-border-strong: #CDD3DF;
        --ca-text: #12162A;
        --ca-text-secondary: #5B6172;
        --ca-text-muted: #8A90A3;
        --ca-brand: #16213E;
        --ca-brand-2: #24345C;
        --ca-accent: #0F9E96;
        --ca-accent-text: #0B7A73;
        --ca-accent-light: #E3F5F3;
        --ca-success: #15803D;
        --ca-success-bg: #E7F7ED;
        --ca-danger: #B91C1C;
        --ca-danger-bg: #FDECEC;
        --ca-radius: 12px;
        --ca-radius-sm: 8px;
        --ca-shadow: 0 1px 2px rgba(16,24,40,.04), 0 2px 8px rgba(16,24,40,.06);
    }

    html, body, [class*="css"] {
        font-family: -apple-system, "Segoe UI", Inter, Roboto, Helvetica, Arial, sans-serif;
    }
    .stApp { background: var(--ca-bg); }

    /* Hide default Streamlit chrome for a custom app feel, but keep the
       sidebar re-expand control (it lives inside stToolbar) usable. */
    [data-testid="stToolbarActions"] { visibility: hidden; }
    [data-testid="stToolbar"] { height: 0; }
    [data-testid="stExpandSidebarButton"] { visibility: visible !important; }
    footer { visibility: hidden; height: 0; }
    [data-testid="stAppViewBlockContainer"], .block-container {
        max-width: 980px;
        padding-top: 1.75rem;
        padding-bottom: 2rem;
    }

    /* ---------- Header ---------- */
    .ca-header {
        display: flex;
        align-items: center;
        justify-content: space-between;
        gap: 1rem;
        flex-wrap: wrap;
        padding: .6rem 0 1.1rem 0;
        margin-bottom: 1.1rem;
        border-bottom: 1px solid var(--ca-border);
        position: sticky;
        top: 0;
        z-index: 5;
        background: var(--ca-bg);
    }
    .ca-header-left { display: flex; align-items: center; gap: .85rem; }
    .ca-brand-mark {
        width: 42px; height: 42px; border-radius: 10px;
        background: linear-gradient(135deg, var(--ca-brand) 0%, var(--ca-brand-2) 100%);
        color: #fff; display: flex; align-items: center; justify-content: center;
        font-weight: 700; font-size: .92rem; letter-spacing: .02em;
        box-shadow: var(--ca-shadow);
        flex-shrink: 0;
    }
    .ca-brand-mark.has-logo, .ca-side-brand-mark.has-logo {
        background: transparent; box-shadow: none;
    }
    .ca-brand-mark img, .ca-side-brand-mark img {
        width: 100%; height: 100%; object-fit: contain; display: block;
    }
    .ca-title { font-size: 1.32rem; font-weight: 700; color: var(--ca-text); line-height: 1.2; }
    .ca-tagline { font-size: .84rem; color: var(--ca-text-secondary); margin-top: 2px; }
    .ca-header-right { display: flex; align-items: center; gap: .5rem; flex-wrap: wrap; }

    .ca-pill {
        display: inline-flex; align-items: center; gap: 6px;
        padding: 5px 12px; border-radius: 999px;
        font-size: .76rem; font-weight: 600;
        background: var(--ca-surface-alt); color: var(--ca-text-secondary);
        border: 1px solid var(--ca-border);
        white-space: nowrap;
    }
    .ca-pill.online { background: var(--ca-success-bg); color: var(--ca-success); border-color: transparent; }
    .ca-pill.offline { background: var(--ca-danger-bg); color: var(--ca-danger); border-color: transparent; }
    .ca-dot { width: 7px; height: 7px; border-radius: 50%; flex-shrink: 0; }
    .ca-pill.online .ca-dot { background: #22C55E; animation: ca-pulse 2s infinite; }
    .ca-pill.offline .ca-dot { background: #EF4444; }
    @keyframes ca-pulse {
        0% { box-shadow: 0 0 0 0 rgba(34,197,94,.45); }
        70% { box-shadow: 0 0 0 6px rgba(34,197,94,0); }
        100% { box-shadow: 0 0 0 0 rgba(34,197,94,0); }
    }

    /* ---------- Sidebar ---------- */
    [data-testid="stSidebar"] {
        background: var(--ca-surface);
        border-right: 1px solid var(--ca-border);
    }
    [data-testid="stSidebar"] > div { padding-top: 1.25rem; }
    .ca-side-brand {
        display: flex; align-items: center; gap: .6rem;
        padding-bottom: 1rem; margin-bottom: .9rem;
        border-bottom: 1px solid var(--ca-border);
    }
    .ca-side-brand-mark {
        width: 30px; height: 30px; border-radius: 8px;
        background: linear-gradient(135deg, var(--ca-brand) 0%, var(--ca-brand-2) 100%);
        color: #fff; display: flex; align-items: center; justify-content: center;
        font-weight: 700; font-size: .72rem; flex-shrink: 0;
    }
    .ca-side-brand-name { font-weight: 700; font-size: .92rem; color: var(--ca-text); }

    .ca-eyebrow {
        font-size: .68rem; font-weight: 700; letter-spacing: .08em;
        text-transform: uppercase; color: var(--ca-text-muted);
        margin: .2rem 0 .6rem 0;
    }
    .ca-divider { height: 1px; background: var(--ca-border); margin: 1.1rem 0; }

    .ca-doc-card {
        display: flex; align-items: flex-start; gap: 10px;
        padding: 10px 12px; margin-bottom: 8px;
        background: var(--ca-surface-alt);
        border: 1px solid var(--ca-border);
        border-radius: var(--ca-radius-sm);
        transition: border-color .15s ease;
    }
    .ca-doc-card:hover { border-color: var(--ca-border-strong); }
    .ca-doc-icon {
        width: 26px; height: 26px; border-radius: 6px; flex-shrink: 0;
        background: var(--ca-accent-light); color: var(--ca-accent-text);
        display: flex; align-items: center; justify-content: center;
        font-size: .68rem; font-weight: 700;
    }
    .ca-doc-name { font-size: .82rem; font-weight: 600; color: var(--ca-text); line-height: 1.3; word-break: break-word; }
    .ca-doc-count { font-size: .72rem; color: var(--ca-text-muted); margin-top: 1px; }

    .ca-kv-row {
        display: flex; justify-content: space-between; align-items: center;
        padding: 6px 2px; font-size: .8rem;
    }
    .ca-kv-label { color: var(--ca-text-secondary); }
    .ca-kv-value { color: var(--ca-text); font-weight: 600; text-align: right; }

    .ca-side-footer {
        font-size: .72rem; color: var(--ca-text-muted);
        line-height: 1.5; margin-top: 1.2rem;
    }

    /* ---------- Buttons ---------- */
    [data-testid="stButton"] button {
        border-radius: var(--ca-radius-sm) !important;
        border: 1px solid var(--ca-border-strong) !important;
        font-weight: 600 !important;
        transition: all .15s ease !important;
    }
    [data-testid="stButton"] button:hover {
        border-color: var(--ca-accent) !important;
        color: var(--ca-accent-text) !important;
    }

    /* ---------- Empty / welcome states ---------- */
    .ca-card {
        background: var(--ca-surface);
        border: 1px solid var(--ca-border);
        border-radius: var(--ca-radius);
        padding: 1.4rem 1.5rem;
        box-shadow: var(--ca-shadow);
    }
    .ca-empty-icon {
        width: 48px; height: 48px; border-radius: 12px;
        background: var(--ca-surface-alt); color: var(--ca-text-secondary);
        display: flex; align-items: center; justify-content: center;
        font-size: 1.3rem; margin-bottom: .9rem;
        font-family: "Material Symbols Rounded", sans-serif;
    }
    .ca-icon { font-family: "Material Symbols Rounded", sans-serif; font-size: 1.05rem; vertical-align: middle; }
    .ca-empty-title { font-size: 1.05rem; font-weight: 700; color: var(--ca-text); margin-bottom: .35rem; }
    .ca-empty-text { font-size: .87rem; color: var(--ca-text-secondary); line-height: 1.55; }
    .ca-code {
        display: inline-block; margin-top: .6rem;
        background: var(--ca-brand); color: #E8ECF7;
        padding: 8px 12px; border-radius: 8px;
        font-family: "Consolas", "SFMono-Regular", Menlo, monospace;
        font-size: .8rem;
    }
    .ca-welcome-title { font-size: 1.02rem; font-weight: 700; color: var(--ca-text); margin-bottom: .3rem; }
    .ca-welcome-text { font-size: .85rem; color: var(--ca-text-secondary); margin-bottom: .9rem; }

    /* ---------- Chat messages ---------- */
    div[data-testid="stChatMessage"] { margin-bottom: .35rem; }
    div[data-testid="stChatMessageContent"] {
        background: var(--ca-surface);
        border: 1px solid var(--ca-border);
        border-radius: var(--ca-radius);
        padding: .15rem .95rem !important;
        box-shadow: var(--ca-shadow);
        max-width: 78%;
    }
    div[data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatarUser"]) {
        flex-direction: row-reverse;
    }
    div[data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatarUser"]) div[data-testid="stChatMessageContent"] {
        background: var(--ca-brand);
        border-color: var(--ca-brand);
        color: #F3F5FA;
        margin-left: auto;
    }
    div[data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatarUser"]) div[data-testid="stChatMessageContent"] p {
        color: #F3F5FA;
    }
    div[data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatarAssistant"]) div[data-testid="stChatMessageContent"] {
        border-left: 3px solid var(--ca-accent);
    }
    [data-testid="stChatMessageAvatarUser"], [data-testid="stChatMessageAvatarAssistant"] {
        background: var(--ca-surface-alt) !important;
        border: 1px solid var(--ca-border);
    }

    /* ---------- Source citation cards ---------- */
    .ca-source-card {
        display: flex; gap: 10px;
        padding: 10px 0; border-top: 1px solid var(--ca-border);
    }
    .ca-source-card:first-child { border-top: none; padding-top: 4px; }
    .ca-source-rank {
        width: 22px; height: 22px; border-radius: 50%; flex-shrink: 0;
        background: var(--ca-surface-alt); color: var(--ca-text-secondary);
        display: flex; align-items: center; justify-content: center;
        font-size: .68rem; font-weight: 700; margin-top: 1px;
    }
    .ca-source-body { flex: 1; min-width: 0; }
    .ca-source-meta { display: flex; align-items: baseline; gap: 8px; flex-wrap: wrap; margin-bottom: 4px; }
    .ca-source-file { font-size: .82rem; font-weight: 700; color: var(--ca-text); }
    .ca-source-chunk {
        font-size: .68rem; font-weight: 600; color: var(--ca-text-muted);
        background: var(--ca-surface-alt); border-radius: 999px; padding: 1px 8px;
    }
    .ca-relevance-track {
        height: 4px; border-radius: 999px; background: var(--ca-surface-alt);
        margin: 5px 0 6px 0; overflow: hidden;
    }
    .ca-relevance-fill { height: 100%; background: var(--ca-accent); border-radius: 999px; }
    .ca-source-excerpt {
        font-size: .78rem; color: var(--ca-text-secondary); line-height: 1.5;
        background: var(--ca-surface-alt); border-radius: 8px; padding: 8px 10px;
    }
    .ca-source-distance { font-size: .68rem; color: var(--ca-text-muted); margin-top: 4px; }

    /* ---------- Guardrail / error banners ---------- */
    .ca-banner {
        display: flex; gap: 10px; align-items: flex-start;
        padding: 10px 12px; border-radius: 8px;
        font-size: .84rem; line-height: 1.5;
    }
    .ca-banner.warning { background: #FEF3E2; color: #92400E; }
    .ca-banner.error { background: var(--ca-danger-bg); color: var(--ca-danger); }
    div[data-testid="stChatMessage"]:has(.ca-banner) div[data-testid="stChatMessageContent"] {
        border-left: 3px solid #D97706;
    }

    /* ---------- Attachment chips ---------- */
    .ca-attach-row { display: flex; flex-wrap: wrap; gap: 6px; margin: 2px 0 6px 0; }
    .ca-attach-chip {
        display: inline-flex; align-items: center; gap: 6px;
        padding: 4px 10px 4px 8px; border-radius: 999px;
        font-size: .74rem; font-weight: 600;
        background: var(--ca-surface-alt); color: var(--ca-text-secondary);
        border: 1px solid var(--ca-border);
        max-width: 100%; white-space: nowrap; overflow: hidden; text-overflow: ellipsis;
    }
    .ca-attach-ico { font-size: .9rem; line-height: 1; }
    /* On the dark user bubble, make chips readable */
    div[data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatarUser"]) .ca-attach-chip {
        background: rgba(255,255,255,.14); color: #EAEEF7; border-color: rgba(255,255,255,.22);
    }

    /* ---------- Chat input ---------- */
    [data-testid="stChatInput"] {
        border-radius: var(--ca-radius) !important;
        border: 1px solid var(--ca-border-strong) !important;
        box-shadow: var(--ca-shadow) !important;
    }
    [data-testid="stChatInput"]:focus-within {
        border-color: var(--ca-accent) !important;
    }
    .ca-disclaimer {
        font-size: .72rem; color: var(--ca-text-muted);
        text-align: center; margin-top: .5rem;
    }
    </style>
    """, unsafe_allow_html=True)


# --------------------------------------------------------------------------
# HTML fragment builders
# --------------------------------------------------------------------------
def brand_mark_html(css_class):
    logo_b64 = load_logo_base64()
    if logo_b64:
        return '<div class="{cls} has-logo"><img src="data:image/png;base64,{b64}" alt="{name} logo"/></div>'.format(
            cls=css_class, b64=logo_b64, name=esc(APP_NAME))
    return '<div class="{cls}">CA</div>'.format(cls=css_class)


def attach_icon(name):
    ext = file_ext(name)
    if ext in PDF_EXTS:
        return "📄"
    if ext in IMAGE_EXTS:
        return "🖼️"
    if ext in DOCX_EXTS or ext in CONVERT_EXTS:
        return "📝"
    return "📎"


def attachment_chips_html(names):
    chips = "".join(
        '<span class="ca-attach-chip"><span class="ca-attach-ico">{ico}</span>{name}</span>'.format(
            ico=attach_icon(n), name=esc(n))
        for n in names
    )
    return '<div class="ca-attach-row">{}</div>'.format(chips)


def doc_card_html(name, count):
    initials = "".join(w[0] for w in re.findall(r"[A-Za-z0-9]+", name)[:2]).upper() or "DOC"
    return """
    <div class="ca-doc-card">
        <div class="ca-doc-icon">{initials}</div>
        <div>
            <div class="ca-doc-name">{name}</div>
            <div class="ca-doc-count">{count} chunks indexed</div>
        </div>
    </div>
    """.format(initials=esc(initials), name=esc(name), count=count)


def kv_row_html(label, value):
    return """
    <div class="ca-kv-row">
        <span class="ca-kv-label">{label}</span>
        <span class="ca-kv-value">{value}</span>
    </div>
    """.format(label=esc(label), value=esc(str(value)))


def source_card_html(rank, filename, chunk_id, distance, relevance_pct, excerpt):
    return """
    <div class="ca-source-card">
        <div class="ca-source-rank">{rank}</div>
        <div class="ca-source-body">
            <div class="ca-source-meta">
                <span class="ca-source-file">{filename}</span>
                <span class="ca-source-chunk">chunk {chunk_id}</span>
            </div>
            <div class="ca-relevance-track"><div class="ca-relevance-fill" style="width:{relevance_pct}%"></div></div>
            <div class="ca-source-excerpt">{excerpt}</div>
            <div class="ca-source-distance">L2 distance: {distance:.4f} (lower is closer)</div>
        </div>
    </div>
    """.format(
        rank=rank, filename=esc(filename), chunk_id=chunk_id,
        relevance_pct=relevance_pct, excerpt=esc(excerpt), distance=distance,
    )


# --------------------------------------------------------------------------
# Layout sections
# --------------------------------------------------------------------------
def render_header(doc_count, chunk_count, ollama_online, model_ready):
    if not ollama_online:
        status_class, status_text = "offline", "Model offline"
    elif not model_ready:
        status_class, status_text = "offline", "Model not pulled"
    else:
        status_class, status_text = "online", "Model online"

    st.markdown("""
    <div class="ca-header">
        <div class="ca-header-left">
            {brand_mark}
            <div>
                <div class="ca-title">{app_name}</div>
                <div class="ca-tagline">{tagline}</div>
            </div>
        </div>
        <div class="ca-header-right">
            <div class="ca-pill">{doc_count} document{doc_plural} &middot; {chunk_count} chunks</div>
            <div class="ca-pill {status_class}"><span class="ca-dot"></span>{status_text}</div>
        </div>
    </div>
    """.format(
        brand_mark=brand_mark_html("ca-brand-mark"),
        app_name=esc(APP_NAME), tagline=esc(APP_TAGLINE),
        doc_count=doc_count, doc_plural="" if doc_count == 1 else "s",
        chunk_count=chunk_count, status_class=status_class, status_text=status_text,
    ), unsafe_allow_html=True)


# --------------------------------------------------------------------------
# Legal letterhead PDF export
# --------------------------------------------------------------------------
def get_advocate():
    return {
        "name": (st.session_state.get("adv_name") or "Adv. Ankur Gaurav").strip(),
        "enrolment": (st.session_state.get("adv_enrol") or "").strip(),
        "place": (st.session_state.get("adv_place") or "").strip(),
    }


def _sources_key(sources):
    slim = [
        {
            "source": s.get("source", ""),
            "chunk_id": s.get("chunk_id", ""),
            "text": s.get("text", ""),
        }
        for s in (sources or [])
    ]
    return json.dumps(slim, ensure_ascii=False, sort_keys=True)


@st.cache_data(show_spinner=False)
def _answer_pdf_bytes(question, answer, sources_key, adv_name, adv_enrol, adv_place):
    entry = {"question": question, "answer": answer, "sources": json.loads(sources_key)}
    adv = {"name": adv_name, "enrolment": adv_enrol, "place": adv_place}
    return letterhead.build_legal_pdf([entry], advocate=adv)


@st.cache_data(show_spinner=False)
def _session_pdf_bytes(entries_key, adv_name, adv_enrol, adv_place):
    entries = json.loads(entries_key)
    adv = {"name": adv_name, "enrolment": adv_enrol, "place": adv_place}
    return letterhead.build_legal_pdf(entries, advocate=adv)


def render_answer_download(question, answer, sources, key):
    if not answer:
        return
    adv = get_advocate()
    try:
        pdf_bytes = _answer_pdf_bytes(
            question or "", answer, _sources_key(sources),
            adv["name"], adv["enrolment"], adv["place"])
    except Exception as exc:  # noqa: BLE001
        st.caption("Letterhead PDF unavailable: {}".format(exc))
        return
    fname = "Compliance-Opinion-{}.pdf".format(datetime.date.today().isoformat())
    st.download_button(
        "Download as legal opinion (PDF)", data=pdf_bytes, file_name=fname,
        mime="application/pdf", key=key, icon=":material/verified:")


def render_sidebar(chunks, ollama_online, model_ready):
    with st.sidebar:
        st.markdown("""
        <div class="ca-side-brand">
            {brand_mark}
            <div class="ca-side-brand-name">{app_name}</div>
        </div>
        """.format(
            brand_mark=brand_mark_html("ca-side-brand-mark"), app_name=esc(APP_NAME),
        ), unsafe_allow_html=True)

        st.markdown('<div class="ca-eyebrow">Knowledge Base</div>', unsafe_allow_html=True)
        if not chunks:
            st.markdown(
                '<div class="ca-kv-row"><span class="ca-kv-label">No documents indexed yet</span></div>',
                unsafe_allow_html=True,
            )
        else:
            sources = sorted(set(c["source"] for c in chunks))
            cards = "".join(
                doc_card_html(os.path.basename(src), sum(1 for c in chunks if c["source"] == src))
                for src in sources
            )
            st.markdown(cards, unsafe_allow_html=True)

        st.markdown('<div class="ca-divider"></div>', unsafe_allow_html=True)
        st.markdown('<div class="ca-eyebrow">System Status</div>', unsafe_allow_html=True)
        rows = (
            kv_row_html("Language model", OLLAMA_MODEL)
            + kv_row_html("Embedding model", EMBEDDING_MODEL)
            + kv_row_html("Retrieval depth", "Top {}".format(TOP_K))
            + kv_row_html("Ollama connection", "Online" if ollama_online else "Offline")
        )
        st.markdown(rows, unsafe_allow_html=True)
        if ollama_online and not model_ready:
            st.markdown(
                '<div class="ca-banner warning" style="margin-top:.5rem;">'
                'Model "{}" not found. Run <code>ollama pull {}</code>.</div>'.format(
                    esc(OLLAMA_MODEL), esc(OLLAMA_MODEL)
                ),
                unsafe_allow_html=True,
            )

        attachments = st.session_state.get("attachments", [])
        if attachments:
            st.markdown('<div class="ca-divider"></div>', unsafe_allow_html=True)
            st.markdown('<div class="ca-eyebrow">Session Attachments</div>', unsafe_allow_html=True)
            cards = "".join(
                doc_card_html(att["name"], len(att["chunks"])) for att in attachments
            )
            st.markdown(cards, unsafe_allow_html=True)

        st.markdown('<div class="ca-divider"></div>', unsafe_allow_html=True)
        st.markdown('<div class="ca-eyebrow">Legal Letterhead</div>', unsafe_allow_html=True)
        with st.expander("Advocate & document details", icon=":material/stylus_note:"):
            st.text_input("Advocate name", key="adv_name")
            st.text_input("Enrolment no.", key="adv_enrol", placeholder="e.g. D/1234/2019")
            st.text_input("Place", key="adv_place", placeholder="e.g. New Delhi")
            st.caption("Used in the header, footer verification and signature block "
                       "of the exported PDF.")

        answered = [
            {"question": m.get("question", ""), "answer": m["content"],
             "sources": m.get("sources", [])}
            for m in st.session_state.get("messages", [])
            if m["role"] == "assistant" and m.get("kind", "normal") == "normal" and m["content"]
        ]
        if answered:
            adv = get_advocate()
            try:
                full_pdf = _session_pdf_bytes(
                    json.dumps(answered, ensure_ascii=False, sort_keys=True),
                    adv["name"], adv["enrolment"], adv["place"])
                st.download_button(
                    "Download full session (PDF)", data=full_pdf,
                    file_name="Compliance-Report-{}.pdf".format(datetime.date.today().isoformat()),
                    mime="application/pdf", use_container_width=True,
                    icon=":material/verified:")
            except Exception as exc:  # noqa: BLE001
                st.caption("Report export unavailable: {}".format(exc))

        st.markdown('<div class="ca-divider"></div>', unsafe_allow_html=True)
        st.markdown('<div class="ca-eyebrow">Session</div>', unsafe_allow_html=True)
        if st.button("Clear conversation", use_container_width=True, icon=":material/delete_sweep:"):
            st.session_state.messages = []
            st.session_state.attachments = []
            st.rerun()

        st.markdown(
            '<div class="ca-side-footer">Answers are generated by a local LLM from your indexed '
            'documents. Always verify against source material before relying on this for compliance '
            'decisions.</div>',
            unsafe_allow_html=True,
        )


def render_empty_state():
    st.markdown("""
    <div class="ca-card">
        <div class="ca-empty-icon">description</div>
        <div class="ca-empty-title">No vector database found</div>
        <div class="ca-empty-text">
            Index your compliance documents before starting a conversation. Add PDFs to the
            <code>data/</code> folder, then build the vector database:
            <div class="ca-code">python create_vector_db.py</div>
        </div>
    </div>
    """, unsafe_allow_html=True)


def render_welcome():
    st.markdown("""
    <div class="ca-card" style="margin-bottom:1rem;">
        <div class="ca-welcome-title">Ask about your compliance documents</div>
        <div class="ca-welcome-text">
            Retrieval-augmented answers grounded in your indexed PDFs, with source citations for every response.
            Attach a PDF, screenshot or Word file to the message box to ask about it directly.
        </div>
    </div>
    """, unsafe_allow_html=True)
    cols = st.columns(len(EXAMPLE_QUESTIONS))
    for col, question in zip(cols, EXAMPLE_QUESTIONS):
        with col:
            if st.button(question, use_container_width=True, key="example_{}".format(hash(question))):
                st.session_state.pending_question = question
                st.rerun()


def render_sources(sources):
    if not sources:
        return
    with st.expander("Sources ({})".format(len(sources)), icon=":material/link:"):
        distances = [s["score"] for s in sources]
        lo, hi = min(distances), max(distances)
        span = (hi - lo) or 1.0
        cards = ""
        for i, s in enumerate(sources):
            relevance_pct = max(12, round((1 - (s["score"] - lo) / span) * 100))
            normalized_text = re.sub(r"\s+", " ", s["text"]).strip()
            excerpt = normalized_text[:220] + ("…" if len(normalized_text) > 220 else "")
            cards += source_card_html(
                rank=i + 1,
                filename=os.path.basename(s["source"]),
                chunk_id=s["chunk_id"],
                distance=s["score"],
                relevance_pct=relevance_pct,
                excerpt=excerpt,
            )
        st.markdown(cards, unsafe_allow_html=True)


def render_message(msg, idx=0):
    role = msg["role"]
    avatar = ":material/gavel:" if role == "assistant" else None
    with st.chat_message(role, avatar=avatar):
        if msg.get("attachments"):
            st.markdown(attachment_chips_html(msg["attachments"]), unsafe_allow_html=True)
        kind = msg.get("kind", "normal")
        if kind == "blocked":
            st.markdown(
                '<div class="ca-banner warning"><span class="ca-icon">shield</span>'
                '<span>{}</span></div>'.format(esc(msg["content"])),
                unsafe_allow_html=True,
            )
        elif kind == "error":
            st.markdown(
                '<div class="ca-banner error"><span class="ca-icon">error</span>'
                '<span>{}</span></div>'.format(esc(msg["content"])),
                unsafe_allow_html=True,
            )
        else:
            if msg["content"]:
                st.markdown(msg["content"])
            if msg.get("sources"):
                render_sources(msg["sources"])
            if role == "assistant" and msg["content"]:
                render_answer_download(
                    msg.get("question", ""), msg["content"], msg.get("sources"),
                    key="dl_hist_{}".format(idx))


def handle_user_message(text, files, index, chunks, embedding_model, chat_container):
    with chat_container:
        attachment_names = [f.name for f in files]

        with st.chat_message("user"):
            if attachment_names:
                st.markdown(attachment_chips_html(attachment_names), unsafe_allow_html=True)
            if text:
                st.markdown(text)
        st.session_state.messages.append({
            "role": "user", "content": text, "kind": "normal",
            "attachments": attachment_names,
        })

        # Extract, chunk and embed any newly attached files.
        if files:
            with st.status("Reading {} attachment{}…".format(
                    len(files), "" if len(files) == 1 else "s"), expanded=True) as status:
                notes = process_attachments(files, embedding_model)
                icons = {"ok": "✅", "empty": "•", "error": "⚠️", "duplicate": "↺"}
                for name, state, detail in notes:
                    st.write("{} **{}** — {}".format(icons.get(state, "•"), esc(name), detail))
                status.update(label="Attachments processed", state="complete", expanded=False)

        # Question sent to the model; default to a summary ask when only files were sent.
        question = text or DEFAULT_ATTACH_PROMPT

        blocked, term = check_guardrails(question)
        if blocked:
            content = GUARDRAIL_MESSAGE.format(term)
            with st.chat_message("assistant", avatar=":material/gavel:"):
                st.markdown(
                    '<div class="ca-banner warning"><span class="ca-icon">shield</span>'
                    '<span>{}</span></div>'.format(esc(content)),
                    unsafe_allow_html=True,
                )
            st.session_state.messages.append({"role": "assistant", "content": content, "kind": "blocked"})
            return

        query_vec = embed_query(embedding_model, question)
        attach_results = search_attachments(query_vec, ATTACH_TOP_K)
        kb_results = search_kb(query_vec, index, chunks, TOP_K)
        # Attachments first so freshly provided context is prioritised in the prompt.
        retrieved = attach_results + kb_results

        context_text = ""
        for i, r in enumerate(retrieved):
            context_text += "[{}] {} (chunk {}):\n{}\n\n".format(
                i + 1, os.path.basename(r["source"]), r["chunk_id"], r["text"])
        history_for_ollama = [
            {"role": m["role"], "content": m["content"]}
            for m in st.session_state.messages[:-1]
            if m.get("kind", "normal") == "normal" and m["content"]
        ]
        messages = build_prompt_messages(history_for_ollama, context_text, question)

        with st.chat_message("assistant", avatar=":material/gavel:"):
            placeholder = st.empty()
            placeholder.markdown('<span style="color:var(--ca-text-muted);">Generating response…</span>',
                                  unsafe_allow_html=True)
            full_response = ""
            for token in stream_ollama(messages):
                full_response += token
                placeholder.markdown(full_response + "▌")
            placeholder.markdown(full_response)

            if full_response.startswith("ERROR:"):
                placeholder.markdown(
                    '<div class="ca-banner error"><span class="ca-icon">error</span>'
                    '<span>{}</span></div>'.format(esc(full_response)),
                    unsafe_allow_html=True,
                )
                st.session_state.messages.append(
                    {"role": "assistant", "content": full_response, "kind": "error"})
                return

            if retrieved:
                render_sources(retrieved)
            render_answer_download(
                question, full_response, retrieved,
                key="dl_live_{}".format(len(st.session_state.messages)))
        st.session_state.messages.append({
            "role": "assistant", "content": full_response, "sources": retrieved,
            "question": question, "kind": "normal"})


# --------------------------------------------------------------------------
# Entry point
# --------------------------------------------------------------------------
def parse_user_input(raw):
    """Normalise st.chat_input output (string, ChatInputValue or None).

    Returns (text, files) where text is a stripped string and files is a list
    of uploaded files (possibly empty).
    """
    if raw is None:
        return "", []
    if isinstance(raw, str):
        return raw.strip(), []
    text = (getattr(raw, "text", "") or "").strip()
    files = list(getattr(raw, "files", []) or [])
    return text, files


def main():
    page_icon = FAVICON_FILE if os.path.exists(FAVICON_FILE) else "⚖️"
    st.set_page_config(page_title=APP_NAME, page_icon=page_icon, layout="wide")
    inject_css()

    st.session_state.setdefault("adv_name", "Adv. Ankur Gaurav")
    st.session_state.setdefault("adv_enrol", "")
    st.session_state.setdefault("adv_place", "")

    index, chunks = load_vectorstore()
    embedding_model = load_embedding_model()
    ollama_online, model_ready = check_ollama_status()
    doc_count = len(set(c["source"] for c in chunks)) if chunks else 0
    chunk_count = len(chunks) if chunks else 0

    render_sidebar(chunks, ollama_online, model_ready)

    page = st.container()
    with page:
        render_header(doc_count, chunk_count, ollama_online, model_ready)

        if index is None:
            render_empty_state()
            return

        if "messages" not in st.session_state:
            st.session_state.messages = []

        raw_input = st.chat_input(
            "Ask a compliance question, or attach a PDF, screenshot or document…",
            accept_file="multiple",
            file_type=SUPPORTED_UPLOAD_EXTS,
        )
        if raw_input is None and st.session_state.get("pending_question"):
            raw_input = st.session_state.pop("pending_question")
        user_text, user_files = parse_user_input(raw_input)
        has_input = bool(user_text or user_files)

        if not st.session_state.messages and not has_input:
            render_welcome()
            chat_container = None
        else:
            chat_container = st.container(height=560)
            with chat_container:
                for i, msg in enumerate(st.session_state.messages):
                    render_message(msg, i)

        if has_input:
            handle_user_message(user_text, user_files, index, chunks, embedding_model, chat_container)

        st.markdown(
            '<div class="ca-disclaimer">AI-generated answers may be inaccurate or incomplete. '
            'Verify against source documents before relying on this for compliance decisions.</div>',
            unsafe_allow_html=True,
        )


if __name__ == "__main__":
    main()
